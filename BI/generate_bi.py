import os
import pandas as pd
import plotly.express as px
import plotly.io as pio
from docx import Document
import nbformat as nbf
import io

# 1. Setup Data
base_dir = "/home/mirahasina/L3_INSI/BI/E-sitrana/BI"
out_dir = os.path.join(base_dir, "Resultats_BI")
os.makedirs(out_dir, exist_ok=True)

# Copy data to the new folder
os.system(f'cp "{base_dir}/esitrana_export_2026-05-01.csv" "{out_dir}/"')
os.system(f'cp "{base_dir}/rendez_vous_esitrana.csv" "{out_dir}/"')

# 2. Extract Data
with open(os.path.join(base_dir, 'esitrana_export_2026-05-01.csv'), 'r') as f:
    lines = f.readlines()

data_dict = {}
current_key = None
current_data = []

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line in ["EXPORT DES DONNEES E-SITRANA", "Date d'export: 01/05/2026 11:02:28"]:
        continue
    if line in ["RENDEZ-VOUS", "PATIENTS", "MEDECINS", "SERVICES"]:
        if current_key and current_data:
            data_dict[current_key] = current_data
        current_key = line
        current_data = []
    else:
        current_data.append(line)
if current_key and current_data:
    data_dict[current_key] = current_data

# Create DataFrames
df_rv = pd.read_csv(io.StringIO('\n'.join(data_dict['RENDEZ-VOUS'])))
df_patients = pd.read_csv(io.StringIO('\n'.join(data_dict['PATIENTS'])))
df_medecins = pd.read_csv(io.StringIO('\n'.join(data_dict['MEDECINS'])))
df_services = pd.read_csv(io.StringIO('\n'.join(data_dict['SERVICES'])))

# Transform
df_rv['Date'] = pd.to_datetime(df_rv['Date'])
df_rv['Mois'] = df_rv['Date'].dt.month_name()
df_rv['Jour'] = df_rv['Date'].dt.day_name()

# Handle some potential encoding issues with stats
df_rv['Statut'] = df_rv['Statut'].replace('terminÃ©', 'terminé')

# 3. Create HTML Dashboard
fig1 = px.histogram(df_rv, x='Date', title='Evolution des Rendez-Vous dans le temps', nbins=30)
fig2 = px.pie(df_rv, names='Service', title='Répartition des Rendez-Vous par Service')
fig3 = px.pie(df_rv, names='Statut', title='Statut des Rendez-Vous', color='Statut', color_discrete_map={'terminé':'green', 'en suivi':'blue', 'en attente':'orange'})
fig4 = px.histogram(df_rv, x='Medecin', title='Rendez-vous par Médecin', color='Statut')

html_content = f"""
<html>
<head><title>Tableau de Bord BI - E-Sitrana</title>
<style>
body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 0; padding: 20px; background-color: #f0f2f5; }}
h1 {{ color: #1a237e; text-align: center; padding: 20px 0; }}
.chart-container {{ display: flex; flex-wrap: wrap; justify-content: space-around; }}
.chart {{ width: 45%; min-width: 500px; margin-bottom: 30px; background: white; padding: 15px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }}
</style>
</head>
<body>
<h1>Tableau de Bord Décisionnel (Business Intelligence)</h1>
<div class="chart-container">
    <div class="chart">{fig1.to_html(full_html=False, include_plotlyjs='cdn')}</div>
    <div class="chart">{fig2.to_html(full_html=False, include_plotlyjs='cdn')}</div>
    <div class="chart">{fig3.to_html(full_html=False, include_plotlyjs='cdn')}</div>
    <div class="chart">{fig4.to_html(full_html=False, include_plotlyjs='cdn')}</div>
</div>
</body>
</html>
"""
with open(os.path.join(out_dir, "dashboard_bi.html"), "w", encoding="utf-8") as f:
    f.write(html_content)

# 4. Create Word Document (Rapport)
doc = Document()
doc.add_heading('Rapport Business Intelligence - E-Sitrana', 0)

doc.add_heading('1. Architecture du Data Warehouse', level=1)
doc.add_paragraph('Le Data Warehouse a été modélisé en étoile (Star Schema) pour optimiser les requêtes analytiques.')
doc.add_heading('Table des Faits :', level=2)
doc.add_paragraph('- Fait_RendezVous : Mesure l\'activité des consultations. (Clés étrangères : id_patient, id_medecin, id_service, id_temps, attributs : statut, heure).')
doc.add_heading('Dimensions :', level=2)
doc.add_paragraph('- Dim_Patient : Informations sur les patients (âge, localisation, sexe).')
doc.add_paragraph('- Dim_Medecin : Informations sur les médecins (spécialité).')
doc.add_paragraph('- Dim_Service : Informations sur les services (nom, description).')
doc.add_paragraph('- Dim_Temps : Hiérarchie temporelle (Année, Mois, Jour, Jour de la semaine).')

doc.add_heading('2. Processus ETL (Extract, Transform, Load)', level=1)
doc.add_paragraph("Extraction : Les données sont extraites des exports ERP au format CSV (esitrana_export.csv et rendez_vous_esitrana.csv).")
doc.add_paragraph("Transformation : Nettoyage des données (correction d'encodage), conversion des dates, création de la dimension Temps à partir des dates de rendez-vous, résolution des clés.")
doc.add_paragraph("Chargement : Intégration dans des structures de données (DataFrames / Base de données) orientées analyse.")

doc.add_heading('3. Indicateurs de Performance (KPIs)', level=1)
doc.add_paragraph("À partir de l'analyse, nous avons extrait les indicateurs suivants pour suivre l'évolution des activités :")
total_rv = len(df_rv)
rv_termines = len(df_rv[df_rv['Statut'] == 'terminé'])
doc.add_paragraph(f"- Nombre total de rendez-vous enregistrés : {total_rv}")
doc.add_paragraph(f"- Taux d'achèvement des rendez-vous : {rv_termines/total_rv:.2%}")
doc.add_paragraph("- Top 3 des services les plus sollicités : " + ", ".join(df_rv['Service'].value_counts().head(3).index.tolist()))

doc.add_paragraph("\nConclusion : L'activité est concentrée principalement en Médecine générale, avec un fort taux de rendez-vous en cours de suivi. L'ouverture de l'exploration de ces données via le dashboard interactif (dashboard_bi.html) permettra au management de mieux répartir les ressources médicales.")

doc.save(os.path.join(out_dir, "rapport_bi.docx"))

# 5. Create Jupyter Notebook (.ipynb)
nb = nbf.v4.new_notebook()

text = "# Projet BI - E-Sitrana : Data Warehouse, ETL et Visualisation\nCe notebook contient le processus ETL et la création des modèles pour le Data Warehouse (Modélisation en Étoile)."
code_etl = '''import pandas as pd
import io
import matplotlib.pyplot as plt

# 1. EXTRACTION
with open('esitrana_export_2026-05-01.csv', 'r') as f:
    lines = f.readlines()

data_dict = {}
current_key = None
current_data = []
for line in lines:
    line = line.strip()
    if not line or line in ["EXPORT DES DONNEES E-SITRANA", "Date d'export: 01/05/2026 11:02:28"]: continue
    if line in ["RENDEZ-VOUS", "PATIENTS", "MEDECINS", "SERVICES"]:
        if current_key and current_data: data_dict[current_key] = current_data
        current_key = line
        current_data = []
    else: current_data.append(line)
if current_key and current_data: data_dict[current_key] = current_data

df_rv = pd.read_csv(io.StringIO('\\n'.join(data_dict['RENDEZ-VOUS'])))

# 2. TRANSFORMATION (Modélisation en Étoile)
# Dim_Temps
df_rv['Date'] = pd.to_datetime(df_rv['Date'])
dim_temps = pd.DataFrame({'Date': df_rv['Date'].unique()})
dim_temps['Annee'] = dim_temps['Date'].dt.year
dim_temps['Mois'] = dim_temps['Date'].dt.month
dim_temps['Jour_Semaine'] = dim_temps['Date'].dt.day_name()
dim_temps['ID_Temps'] = range(1, len(dim_temps) + 1)

# Fusion pour Table des Faits
fait_rv = df_rv.merge(dim_temps, on='Date', how='left')
fait_rv = fait_rv[['ID', 'Patient', 'Medecin', 'Service', 'ID_Temps', 'Heure', 'Statut']]

# Nettoyage
fait_rv['Statut'] = fait_rv['Statut'].replace('terminÃ©', 'terminé')

# 3. CHARGEMENT ET VISUALISATION
print("Table des faits Fait_RendezVous (Aperçu) :")
display(fait_rv.head())

# Visualisation
fait_rv.groupby('Service').size().plot(kind='bar', title='Nombre de Rendez-vous par Service', color='skyblue')
plt.ylabel('Nombre de RDV')
plt.show()
'''

nb['cells'] = [nbf.v4.new_markdown_cell(text), nbf.v4.new_code_cell(code_etl)]
with open(os.path.join(out_dir, "processus_etl_bi.ipynb"), 'w') as f:
    nbf.write(nb, f)

print("Tous les fichiers ont été générés avec succès dans le dossier Resultats_BI.")
