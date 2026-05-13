from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles globaux ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1E3A8A'):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Calibri'
    return p

def add_para(doc, text, bold=False, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

# ── Page de titre ────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RAPPORT BUSINESS INTELLIGENCE')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
run.font.name = 'Calibri'

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Système d\'Information Médical — E-Sitrana')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
r2.font.name = 'Calibri'

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f'Date d\'analyse : 01 mai 2026 | Généré le {datetime.date.today().strftime("%d/%m/%Y")}')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
r3.font.name = 'Calibri'

doc.add_page_break()

# ── 1. Introduction ──────────────────────────────────────────
add_heading(doc, '1. Introduction et Contexte', level=1)
add_para(doc, "E-Sitrana est un système d'information médical centralisé destiné à la gestion des rendez-vous médicaux, des patients et du personnel soignant. Dans le cadre de l'évolution vers une gestion basée sur les données, la mise en place d'une solution de Business Intelligence (BI) permet de transformer les données opérationnelles en informations décisionnelles exploitables.")
doc.add_paragraph()
add_para(doc, "Ce rapport présente l'ensemble du processus BI, depuis la conception du Data Warehouse jusqu'aux indicateurs de performance, en passant par le processus ETL (Extract, Transform, Load). Les analyses portent sur les données extraites au 01/05/2026.")

# ── 2. Architecture DW ───────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '2. Architecture du Data Warehouse', level=1)
add_heading(doc, '2.1 Modélisation en Étoile (Star Schema)', level=2)
add_para(doc, "Le Data Warehouse adopte une modélisation en étoile, optimisée pour les requêtes analytiques OLAP. Cette architecture se compose d'une table centrale de faits reliée à plusieurs tables de dimensions.")

doc.add_paragraph()
add_para(doc, "Table centrale des faits :", bold=True)
add_para(doc, "• Fait_RendezVous : Enregistre chaque rendez-vous médical. Contient les clés étrangères vers toutes les dimensions ainsi que les mesures : nb_rendez_vous, est_terminé, est_en_suivi, est_en_attente.", indent=True)

doc.add_paragraph()
add_para(doc, "Tables de dimensions :", bold=True)

dims = [
    ("Dim_Temps", "Hiérarchie temporelle complète : Année → Trimestre → Mois → Semaine → Jour → Jour de la semaine"),
    ("Dim_Patient", "Informations sur les patients : identité, date de naissance, âge calculé, localisation géographique"),
    ("Dim_Medecin", "Informations sur les médecins : spécialité, coordonnées"),
    ("Dim_Service", "Catalogue des services médicaux disponibles et leur description"),
]
for name, desc in dims:
    add_para(doc, f"• {name} : {desc}", indent=True)

# ── 3. Processus ETL ─────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '3. Processus ETL', level=1)

add_heading(doc, '3.1 Extraction', level=2)
add_para(doc, "Les données sources sont issues de deux fichiers CSV générés par l'ERP E-Sitrana :")
add_para(doc, "• esitrana_export_2026-05-01.csv : Export multi-sections contenant Rendez-Vous, Patients, Médecins et Services", indent=True)
add_para(doc, "• rendez_vous_esitrana.csv : Vue complémentaire des rendez-vous avec dates en format littéral", indent=True)

add_heading(doc, '3.2 Transformation', level=2)
transformations = [
    "Correction de l'encodage caractères (latin-1 → UTF-8) pour les caractères accentués",
    "Parsing du fichier multi-sections par détection des en-têtes textuels",
    "Conversion des colonnes de dates en type datetime",
    "Génération de la dimension Dim_Temps avec calcul de la semaine ISO, trimestre et libellé du mois",
    "Calcul de l'âge des patients à partir de leur date de naissance",
    "Création des mesures booléennes (est_terminé, est_en_suivi, est_en_attente) dans la table de faits",
    "Normalisation des statuts (suppression des variantes d'encodage)",
]
for t in transformations:
    add_para(doc, f"• {t}", indent=True)

add_heading(doc, '3.3 Chargement', level=2)
add_para(doc, "Les tables résultantes sont exportées en format CSV structuré (dw_fait_rendez_vous.csv, dw_dim_*.csv) et chargées dans l'environnement d'analyse. Dans un contexte de production, ces tables seraient chargées dans un SGBD relationnel (PostgreSQL, MySQL) dédié au Data Warehouse.")

# ── 4. KPIs ─────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '4. Indicateurs Clés de Performance (KPIs)', level=1)
add_heading(doc, '4.1 Vue d\'ensemble', level=2)

# Table KPI globaux
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Indicateur'
hdr[1].text = 'Valeur'
for cell in hdr:
    set_cell_bg(cell, '1E3A8A')
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

kpis = [
    ('Total Rendez-Vous enregistrés', '37'),
    ('Rendez-Vous terminés', '20 (54,1%)'),
    ('Rendez-Vous en suivi', '12 (32,4%)'),
    ('Rendez-Vous en attente', '5 (13,5%)'),
    ('Nombre de patients uniques', '13'),
    ('Nombre de médecins actifs', '11'),
    ('Services médicaux utilisés', '6'),
    ('Période couverte', 'Mars 2026 – Mai 2026'),
]
for label, val in kpis:
    row = table.add_row().cells
    row[0].text = label
    row[1].text = val

doc.add_paragraph()
add_heading(doc, '4.2 Analyse par Service', level=2)

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
headers2 = ['Service', 'Total RDV', 'Terminés', "Taux d'achèvement"]
for i, h in enumerate(headers2):
    cell = t2.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '0E7490')
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

services_data = [
    ('Médecine générale', '20', '13', '65%'),
    ('Pédiatrie', '5', '3', '60%'),
    ('Gynécologie', '4', '2', '50%'),
    ('Cardiologie', '3', '3', '100%'),
    ('Laboratoire', '3', '2', '67%'),
    ('Imagerie Médicale', '2', '1', '50%'),
]
for row_data in services_data:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

# ── 5. Analyse Temporelle ────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '5. Analyse Temporelle et Tendances', level=1)
add_para(doc, "L'analyse de l'évolution mensuelle révèle une forte concentration de l'activité en mars 2026, mois de démarrage du système avec 30 rendez-vous enregistrés. Ce pic s'explique par la mise en production de l'ERP et l'initialisation des données historiques.")
doc.add_paragraph()

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
for i, h in enumerate(['Période', 'Nombre de RDV', 'Observation']):
    cell = t3.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, '1E3A8A')
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

temporal = [
    ('Mars 2026', '30', 'Démarrage système — données historiques'),
    ('Avril 2026', '2', 'Activité courante — faible volume'),
    ('Mai 2026 (partiel)', '5', 'Rendez-vous futurs planifiés'),
]
for row_data in temporal:
    row = t3.add_row().cells
    for i, v in enumerate(row_data):
        row[i].text = v

# ── 6. Recommandations ───────────────────────────────────────
doc.add_paragraph()
add_heading(doc, '6. Conclusions et Recommandations', level=1)

recs = [
    ("Enrichissement des données patients", "Les données de naissance et d'adresse sont manquantes pour les nouveaux patients (IDs 10–13). Il est recommandé de rendre ces champs obligatoires à l'inscription pour enrichir les analyses démographiques."),
    ("Suivi du taux d'achèvement", "Le taux global de 54% indique que près de la moitié des rendez-vous sont encore en cours ou en attente. Un suivi régulier de cet indicateur permettra de détecter les goulots d'étranglement."),
    ("Équilibrage de la charge médicale", "La Médecine générale concentre 54% des rendez-vous. Une analyse plus fine permettrait d'optimiser la répartition entre les médecins généralistes."),
    ("Automatisation ETL", "Mettre en place un pipeline ETL automatisé (Airflow, cron) pour alimenter le Data Warehouse quotidiennement et maintenir les tableaux de bord à jour en temps réel."),
    ("Extension du Data Warehouse", "Intégrer les données de facturation et de prescription médicale pour enrichir l'analyse et calculer des KPIs financiers."),
]

for title_r, desc_r in recs:
    add_para(doc, f"▸ {title_r}", bold=True)
    add_para(doc, desc_r, indent=True)
    doc.add_paragraph()

# ── Footer ───────────────────────────────────────────────────
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run('E-Sitrana BI · Rapport Business Intelligence · Modélisation en Étoile · ETL Python · 2026')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
fr.font.italic = True

doc.save('/home/mirahasina/L3_INSI/BI/E-sitrana/BI/Resultats_BI/rapport_bi.docx')
print("Rapport Word généré avec succès.")
